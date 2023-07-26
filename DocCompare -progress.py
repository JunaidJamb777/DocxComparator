import docx
import spacy
import os 
class DocxComparator:
    def __init__(self, file1_path, file2_path):
        self.file1_path = file1_path
        self.file2_path = file2_path
        self.nlp = spacy.load('en_core_web_sm')

    def read_docx(self, file_path):
        try:
            doc = docx.Document(file_path)
            extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return extracted_text
        except Exception as e:
            print(f"Error occurred while reading the .docx file {file_path}: {e}")
            return None

    def compare_text(self, text1, text2):
        try:
            len1 = len(text1)
            len2 = len(text2)
            dp = [[0] * (len2 + 1) for _ in range(len1 + 1)]
            for i in range(len1 + 1):
                dp[i][0] = i
            for j in range(len2 + 1):
                dp[0][j] = j
            for i in range(1, len1 + 1):
                for j in range(1, len2 + 1):
                    cost = 0 if text1[i - 1] == text2[j - 1] else 1
                    dp[i][j] = min(dp[i - 1][j] + 1, dp[i][j - 1] + 1, dp[i - 1][j - 1] + cost)
            max_length = max(len1, len2)
            similarity_score = 1 - (dp[len1][len2] / max_length)
            return similarity_score
        except Exception as e:
            print("Error occurred while comparing text:", e)
            return None

    def align_text(self, text1, text2):
        try:
            sentences1 = text1.split('\n')
            sentences2 = text2.split('\n')
            len1 = len(sentences1)
            len2 = len(sentences2)
            dp = [[0] * (len2 + 1) for _ in range(len1 + 1)]
            for i in range(len1 + 1):
                dp[i][0] = i
            for j in range(len2 + 1):
                dp[0][j] = j
            for i in range(1, len1 + 1):
                for j in range(1, len2 + 1):
                    cost = 0 if sentences1[i - 1] == sentences2[j - 1] else 1
                    dp[i][j] = min(dp[i - 1][j] + 1, dp[i][j - 1] + 1, dp[i - 1][j - 1] + cost)
            aligned_text = []
            i, j = len1, len2
            while i > 0 and j > 0:
                if sentences1[i - 1] == sentences2[j - 1]:
                    aligned_text.append(sentences1[i - 1])
                    i -= 1
                    j -= 1
                elif dp[i][j] == dp[i - 1][j] + 1:
                    aligned_text.append(sentences1[i - 1])
                    i -= 1
                else:
                    aligned_text.append(sentences2[j - 1])
                    j -= 1
            while i > 0:
                aligned_text.append(sentences1[i - 1])
                i -= 1
            while j > 0:
                aligned_text.append(sentences2[j - 1])
                j -= 1
            aligned_text.reverse()
            return "\n".join(aligned_text)
        except Exception as e:
            print("Error occurred while aligning text:", e)
            return None

    def analyze_and_improve(self, aligned_text):
        try:
            doc = self.nlp(aligned_text)
            updated_sentences = []
            for sentence in doc.sents:
                updated_sentences.append(self.generate_suggestion(sentence))
            return " ".join(updated_sentences)
        except Exception as e:
            print("Error occurred while analyzing and improving text:", e)
            return None

    def generate_suggestion(self, sentence):
        return f"{sentence.text} (Improved)"
    def progress(self, progress_value):
        # Function to display the progress during writing
        print(f"Writing new document: {progress_value:.2f}% complete")
    def generate_updated_docx(self, updated_text):
        output_file = "updated_file.docx"
        total_sentences = len(updated_text.split('\n'))
        updated_doc = docx.Document()  # Create a new Document instance

        try:
            for i, sentence in enumerate(updated_text.split('\n'), 1):
                updated_doc.add_paragraph(sentence)
                progress_value = (i / total_sentences) * 100
                self.progress(progress_value)

            current_dir = os.getcwd()
            output_path = os.path.join(current_dir, output_file)
            updated_doc.save(output_path)
            print(f"\nUpdated .docx File Saved Successfully: {output_path}")
        except Exception as e:
            print(f"Error occurred while saving the updated .docx file: {e}")

# ... (rest of the code)

    def run(self):
        text1 = self.read_docx(self.file1_path)
        text2 = self.read_docx(self.file2_path)

        if text1 is None or text2 is None:
            print("Failed to read .docx files.")
            return

        comparison_result = self.compare_text(text1, text2)
        print(f"Similarity Score: {comparison_result}")

        aligned_text = self.align_text(text1, text2)
        print(f"Aligned Text:\n{aligned_text}")

        updated_text = self.analyze_and_improve(aligned_text)
        print(f"Updated Text:\n{updated_text}")

        self.generate_updated_docx(updated_text)

if __name__ == "__main__":
    print("Please provide the file paths for 'file1.docx' and 'file2.docx' below:")
    file1_path = input("Enter the path to 'file1.docx': ")
    file2_path = input("Enter the path to 'file2.docx': ")

    comparator = DocxComparator(file1_path, file2_path)
    comparator.run()